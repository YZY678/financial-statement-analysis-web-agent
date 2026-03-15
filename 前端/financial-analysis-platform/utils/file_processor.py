import pandas as pd
import numpy as np
import os
import json
from datetime import datetime
import re
from collections import Counter

class FileProcessor:
    """财务报表文件处理器"""
    
    @staticmethod
    def detect_file_encoding(filepath):
        """检测文件编码（不使用第三方库）"""
        # 常见的中文编码列表
        encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'big5', 'utf-16', 'ascii', 'latin-1']
        
        # 尝试读取文件前1024字节来检测编码
        try:
            with open(filepath, 'rb') as f:
                raw_data = f.read(1024)
            
            # 检查UTF-8 BOM
            if raw_data.startswith(b'\\\\xef\\\\xbb\\\\xbf'):
                return 'utf-8-sig'
            
            # 检查UTF-16 BOM
            if raw_data.startswith(b'\\\\xff\\\\xfe') or raw_data.startswith(b'\\\\xfe\\\\xff'):
                return 'utf-16'
            
            # 尝试解码
            for encoding in encodings:
                try:
                    raw_data.decode(encoding)
                    return encoding
                except UnicodeDecodeError:
                    continue
            
        except Exception:
            pass
        
        # 默认返回utf-8
        return 'utf-8'
    
    @staticmethod
    def detect_file_type(df):
        """自动检测财务报表类型"""
        columns = [str(col).lower().replace(' ', '') for col in df.columns]
        
        # 利润表关键词
        income_keywords = ['营业收入', '营业成本', '净利润', '营业利润', '每股收益']
        income_matches = sum(1 for col in columns for kw in income_keywords if kw in col)
        
        # 资产负债表关键词
        balance_keywords = ['资产总计', '负债合计', '所有者权益', '货币资金', '应收账款']
        balance_matches = sum(1 for col in columns for kw in balance_keywords if kw in col)
        
        # 现金流量表关键词
        cashflow_keywords = ['经营活动现金流量', '投资活动现金流量', '筹资活动现金流量']
        cashflow_matches = sum(1 for col in columns for kw in cashflow_keywords if kw in col)
        
        # 科目列内容识别（适配PDF抽取后的“科目/本期数/上年同期数”结构）
        subject_income_matches = 0
        subject_balance_matches = 0
        subject_cashflow_matches = 0

        subject_col = next((col for col in df.columns if str(col).strip() in ['科目', '项目']), None)
        if subject_col is not None:
            subject_values = df[subject_col].astype(str).str.replace(' ', '', regex=False).str.lower()
            income_subject_keywords = ['营业收入', '营业成本', '净利润', '营业利润', '每股收益']
            balance_subject_keywords = ['资产总计', '负债合计', '所有者权益', '货币资金', '应收账款']
            cashflow_subject_keywords = ['经营活动现金流量', '投资活动现金流量', '筹资活动现金流量']

            subject_income_matches = sum(
                subject_values.str.contains(kw, na=False).sum() for kw in income_subject_keywords
            )
            subject_balance_matches = sum(
                subject_values.str.contains(kw, na=False).sum() for kw in balance_subject_keywords
            )
            subject_cashflow_matches = sum(
                subject_values.str.contains(kw, na=False).sum() for kw in cashflow_subject_keywords
            )

        income_score = income_matches + subject_income_matches
        balance_score = balance_matches + subject_balance_matches
        cashflow_score = cashflow_matches + subject_cashflow_matches

        if income_score >= 3 and income_score >= balance_score and income_score >= cashflow_score:
            return 'income_statement'
        elif balance_score >= 3 and balance_score >= cashflow_score:
            return 'balance_sheet'
        elif cashflow_score >= 3:
            return 'cash_flow'
        else:
            return 'unknown'
    
    @staticmethod
    def read_financial_file(filepath, filename):
        """读取财务报表文件"""
        _, ext = os.path.splitext(filename)
        ext = ext.lower()[1:]  # 去掉点号
        
        try:
            if ext == 'csv':
                # 使用内置编码检测
                encoding = FileProcessor.detect_file_encoding(filepath)
                
                print(f"检测到文件编码: {encoding}")
                
                # 尝试不同的读取方式
                try:
                    df = pd.read_csv(filepath, encoding=encoding)
                except Exception as e1:
                    print(f"使用编码 {encoding} 读取失败: {e1}")
                    # 尝试其他常见编码
                    other_encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1', 'iso-8859-1']
                    for enc in other_encodings:
                        if enc != encoding:
                            try:
                                df = pd.read_csv(filepath, encoding=enc)
                                print(f"使用备用编码 {enc} 读取成功")
                                break
                            except:
                                continue
                    else:
                        # 所有编码都失败，尝试不指定编码
                        try:
                            df = pd.read_csv(filepath)
                        except:
                            # 尝试使用分号或制表符分隔
                            try:
                                df = pd.read_csv(filepath, sep=';')
                            except:
                                df = pd.read_csv(filepath, sep='\\\\t')
                
            elif ext in ['xlsx', 'xls']:
                df = pd.read_excel(filepath)

            elif ext == 'docx':
                df = FileProcessor.read_docx_file(filepath)

            elif ext == 'pdf':
                df = FileProcessor.read_pdf_file(filepath)

            elif ext in ['png', 'jpg', 'jpeg']:
                df = FileProcessor.read_image_file(filepath)
                
            else:
                raise ValueError(f"不支持的文件格式: {ext}")
            
            # 清理数据
            df = FileProcessor.clean_financial_data(df)
            
            # 检测报表类型
            report_type = FileProcessor.detect_file_type(df)
            
            print(f"检测到报表类型: {report_type}")
            print(f"数据形状: {df.shape}")
            
            return df, report_type
            
        except Exception as e:
            error_msg = f"读取财务报表失败: {str(e)}"
            print(error_msg)
            raise Exception(error_msg)

    @staticmethod
    def read_docx_file(filepath):
        """读取 DOCX 财务报表"""
        try:
            from docx import Document
        except Exception as e:
            raise Exception("缺少依赖 python-docx，请先安装后再上传 DOCX 文件") from e

        doc = Document(filepath)
        if not doc.tables:
            raise Exception("DOCX 文件中未找到表格")

        # 默认取第一个表格
        table = doc.tables[0]
        rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                rows.append(row_data)

        if not rows:
            raise Exception("DOCX 表格数据为空")

        # 处理表头
        header = rows[0]
        data_rows = rows[1:] if len(rows) > 1 else []

        # 保证列数一致
        max_cols = max(len(r) for r in rows)
        header = (header + [f"列{idx+1}" for idx in range(len(header), max_cols)])[:max_cols]
        normalized_rows = []
        for r in data_rows:
            normalized_rows.append((r + [None] * max(0, max_cols - len(r)))[:max_cols])

        df = pd.DataFrame(normalized_rows, columns=header)
        return df

    @staticmethod
    def read_pdf_file(filepath):
        """读取 PDF 财务报表（优先表格提取，失败回退文本解析）"""
        try:
            import pdfplumber
        except Exception as e:
            raise Exception("缺少依赖 pdfplumber，请先安装后再上传 PDF 文件") from e

        parsed_rows = []

        with pdfplumber.open(filepath) as pdf:
            # 优先提取表格
            for page in pdf.pages:
                tables = page.extract_tables() or []
                for table in tables:
                    for row in table:
                        if not row:
                            continue
                        row_data = [
                            cell.strip() if isinstance(cell, str) else (str(cell).strip() if cell is not None else None)
                            for cell in row
                        ]
                        if any(cell not in [None, ''] for cell in row_data):
                            parsed_rows.append(row_data)

            # 表格为空时，回退文本解析
            if not parsed_rows:
                for page in pdf.pages:
                    text = page.extract_text() or ''
                    lines = [line.strip() for line in text.splitlines() if line and line.strip()]
                    for line in lines:
                        if ',' in line:
                            cols = [c.strip() for c in line.split(',') if c.strip()]
                        elif '\t' in line:
                            cols = [c.strip() for c in line.split('\t') if c.strip()]
                        else:
                            cols = [c for c in re.split(r'\s+', line) if c]
                        if cols:
                            parsed_rows.append(cols)

        if not parsed_rows:
            raise Exception("PDF 中未提取到有效表格或文本")

        # 统一列数
        col_counts = [len(r) for r in parsed_rows if len(r) >= 2]
        if not col_counts:
            raise Exception("PDF 表格列数不足，无法解析")
        target_cols = Counter(col_counts).most_common(1)[0][0]

        normalized = [r for r in parsed_rows if len(r) == target_cols]
        if not normalized:
            normalized = [
                (r + [None] * max(0, target_cols - len(r)))[:target_cols]
                for r in parsed_rows
            ]

        def _is_numeric_cell(value: str) -> bool:
            if value is None:
                return False
            text = str(value).strip().replace(',', '')
            if text == '':
                return False
            return bool(re.fullmatch(r'-?\d+(\.\d+)?%?', text))

        def _row_is_numeric(row):
            return all(_is_numeric_cell(cell) for cell in row)

        header = normalized[0]
        if _row_is_numeric(header):
            header = [f"列{idx+1}" for idx in range(target_cols)]
            data_rows = normalized
        else:
            data_rows = normalized[1:] if len(normalized) > 1 else []

        return pd.DataFrame(data_rows, columns=header)

    @staticmethod
    def read_image_file(filepath):
        """读取图片财务报表（OCR）"""
        try:
            from PIL import Image
        except Exception as e:
            raise Exception("缺少依赖 pillow，请先安装后再上传图片") from e

        try:
            import pytesseract
        except Exception as e:
            raise Exception("缺少依赖 pytesseract，请先安装后再上传图片") from e

        image = Image.open(filepath)

        # 优先尝试中英文识别，失败则回退英文
        try:
            raw_text = pytesseract.image_to_string(image, lang='chi_sim+eng')
        except Exception:
            raw_text = pytesseract.image_to_string(image, lang='eng')

        lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
        if not lines:
            raise Exception("图片中未识别到有效文本")

        # 解析成表格
        parsed_rows = []
        for line in lines:
            if ',' in line:
                cols = [c.strip() for c in line.split(',') if c.strip()]
            elif '\t' in line:
                cols = [c.strip() for c in line.split('\t') if c.strip()]
            else:
                # 优先按任意空白切分，兼容 OCR 单空格输出
                cols = [c for c in re.split(r'\s+', line) if c]
            if cols:
                parsed_rows.append(cols)

        if not parsed_rows:
            raise Exception("图片文本无法解析为表格")

        # 取最常见的列数作为目标列数
        col_counts = [len(r) for r in parsed_rows if len(r) >= 2]
        if not col_counts:
            raise Exception("图片表格列数不足，无法解析")
        target_cols = Counter(col_counts).most_common(1)[0][0]

        normalized = [r for r in parsed_rows if len(r) == target_cols]
        if not normalized:
            # 回退：截断或补齐
            normalized = [
                (r + [None] * max(0, target_cols - len(r)))[:target_cols]
                for r in parsed_rows
            ]

        def _is_numeric_cell(value: str) -> bool:
            if value is None:
                return False
            text = str(value).strip().replace(',', '')
            if text == '':
                return False
            return bool(re.fullmatch(r'-?\d+(\.\d+)?%?', text))

        def _row_is_numeric(row):
            return all(_is_numeric_cell(cell) for cell in row)

        header = normalized[0]
        if _row_is_numeric(header):
            # 首行疑似数据行，构造默认表头
            if target_cols == 4:
                header = ['期间', '营业收入', '营业成本', '净利润']
            elif target_cols == 3:
                header = ['期间', '营业收入', '净利润']
            elif target_cols == 5:
                header = ['期间', '营业收入', '营业成本', '营业利润', '净利润']
            else:
                header = [f"列{idx+1}" for idx in range(target_cols)]
            data_rows = normalized
        else:
            data_rows = normalized[1:] if len(normalized) > 1 else []

        df = pd.DataFrame(data_rows, columns=header)
        return df
    
    @staticmethod
    def clean_financial_data(df):
        """清洗财务数据"""
        print("开始数据清洗...")
        
        # 保存原始列名用于调试
        original_columns = list(df.columns)
        print(f"原始列名: {original_columns}")
        
        # 去除完全重复的行
        original_len = len(df)
        df = df.drop_duplicates()
        if len(df) < original_len:
            print(f"移除了 {original_len - len(df)} 行重复数据")
        
        # 去除全为NaN的列
        original_cols = len(df.columns)
        df = df.dropna(axis=1, how='all')
        if len(df.columns) < original_cols:
            print(f"移除了 {original_cols - len(df.columns)} 个空列")
        
        # 去除全为NaN的行
        original_rows = len(df)
        df = df.dropna(axis=0, how='all')
        if len(df) < original_rows:
            print(f"移除了 {original_rows - len(df)} 行空数据")
        
        # 尝试识别并转换日期列
        date_column = None
        for col in df.columns:
            col_str = str(col)
            # 检查列名是否包含日期相关关键词
            date_keywords = ['期间', '日期', '时间', 'period', 'date', 'year', 'quarter', 'month', 'day']
            if any(keyword in col_str.lower() for keyword in date_keywords):
                date_column = col
                print(f"识别到日期列: {col}")
                try:
                    df[col] = pd.to_datetime(df[col], errors='coerce')
                    print(f"成功转换日期列: {col}")
                except Exception as e:
                    print(f"日期列转换失败: {e}")
                break
        
        # 如果没有找到日期列，尝试第一列
        if date_column is None and len(df.columns) > 0:
            try:
                first_col = df.columns[0]
                parsed_dates = pd.to_datetime(df[first_col], errors='coerce')
                if parsed_dates.notna().any():
                    df[first_col] = parsed_dates
                    date_column = first_col
                    print(f"第一列识别为日期列: {first_col}")
            except:
                pass
        
        # 确保数值列是数值类型
        numeric_cols = []
        for col in df.columns:
            if col != date_column:
                col_name = str(col).strip().lower()
                if col_name in ['科目', '项目', 'subject', 'item']:
                    continue

                # 尝试转换为数值
                try:
                    # 先尝试直接转换
                    parsed_numeric = pd.to_numeric(df[col], errors='coerce')
                    # 如果转换成功，检查是否有数值
                    non_na_count = int(parsed_numeric.notna().sum())
                    if non_na_count > 0:
                        # 仅在可解析数值占比足够高时才覆盖原列，避免破坏文本列
                        if non_na_count / max(len(df), 1) >= 0.4:
                            df[col] = parsed_numeric
                        numeric_cols.append(col)
                except Exception as e:
                    # 转换失败，跳过该列
                    print(f"列 {col} 转换为数值失败: {e}")
        
        print(f"识别到 {len(numeric_cols)} 个数值列: {numeric_cols[:5]}...")
        
        # 排序（如果有日期列）
        if date_column and date_column in df.columns:
            try:
                df = df.sort_values(date_column)
                print(f"按日期列 {date_column} 排序")
            except:
                print("日期排序失败")
        
        print(f"清洗后数据形状: {df.shape}")
        return df
    
    @staticmethod
    def validate_financial_data(df, report_type):
        """验证财务数据的完整性"""
        if df.empty:
            raise ValueError("财务报表为空")
        
        if len(df) < 2:
            raise ValueError("财务报表至少需要2期数据进行分析")
        
        # 检查必要字段
        required_columns = {
            'income_statement': ['营业收入', '净利润'],
            'balance_sheet': ['资产总计', '负债合计', '所有者权益合计'],
            'cash_flow': ['经营活动现金流量净额']
        }
        
        if report_type in required_columns:
            subject_col = next((c for c in df.columns if str(c).strip() in ['科目', '项目']), None)
            subject_values = None
            if subject_col is not None:
                subject_values = df[subject_col].astype(str)

            for col in required_columns[report_type]:
                if col not in df.columns:
                    # 尝试查找相似的列名
                    similar_cols = [c for c in df.columns if col in str(c)]

                    # 兼容长表：在“科目/项目”列中查找必要指标
                    if not similar_cols and subject_values is not None:
                        subject_hit = subject_values.str.contains(col, na=False).any()
                        if subject_hit:
                            print(f"在科目列中找到必要指标: {col}")
                            continue

                    if not similar_cols:
                        # 对于利润表，如果没找到净利润，尝试找利润总额
                        if col == '净利润':
                            alt_cols = ['利润总额', '净收益', '纯利润']
                            similar_cols = [c for c in df.columns if any(alt in str(c) for alt in alt_cols)]
                            if not similar_cols and subject_values is not None:
                                alt_hit = subject_values.str.contains('|'.join(alt_cols), na=False).any()
                                if alt_hit:
                                    print(f"在科目列中找到净利润替代指标: {alt_cols}")
                                    continue
                    
                    if not similar_cols:
                        raise ValueError(f"财务报表缺少必要字段: {col}")
                    else:
                        print(f"找到相似字段 {similar_cols[0]} 代替 {col}")
        
        return True
    
    @staticmethod
    def extract_financial_periods(df):
        """提取财务期间信息"""
        periods = []
        
        # 查找日期列
        date_column = None
        for col in df.columns:
            col_str = str(col)
            date_keywords = ['期间', '日期', '时间', 'period', 'date']
            if any(keyword in col_str.lower() for keyword in date_keywords):
                date_column = col
                break
        
        if date_column and date_column in df.columns:
            for date in df[date_column]:
                try:
                    if isinstance(date, (pd.Timestamp, datetime)):
                        periods.append(date.strftime('%Y-%m-%d'))
                    else:
                        periods.append(str(date))
                except:
                    periods.append(str(date))
        else:
            # 如果没有日期列，使用序号
            periods = [f"第{i+1}期" for i in range(len(df))]
        
        print(f"提取到 {len(periods)} 个期间")
        return periods